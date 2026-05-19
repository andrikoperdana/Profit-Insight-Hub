from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

title = doc.add_heading('Gap Analysis: SecureProfit Hub vs Standard PSO Process', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Dokumen Internal — Untuk Informasi Manajemen')
r.italic = True
r.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Ringkasan Eksekutif', level=1)
doc.add_paragraph(
    "SecureProfit Hub saat ini mencakup sekitar 60% dari standar proses Professional Services Operations (PSO) "
    "4-fase. Kekuatan utama berada pada Fase II (Project Initialization) dan Fase III (Execution & Effort Tracking). "
    "Gap terbesar berada pada Fase I (Lead Lifecycle / CRM) dan Fase IV (Billing & ERP Integration)."
)

def add_phase(title_text, summary_text, table_data, conclusion):
    doc.add_heading(title_text, level=1)
    doc.add_paragraph(summary_text)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Status'
    hdr[2].text = 'Implementasi Saat Ini'
    hdr[3].text = 'Gap'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for row_data in table_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    p = doc.add_paragraph()
    r = p.add_run('Kesimpulan: ')
    r.bold = True
    p.add_run(conclusion)
    doc.add_paragraph()

add_phase(
    'Fase I — Lead Lifecycle & Commercial Design',
    'Fase ini menangani onboarding klien eksternal sebelum project dibuat di sistem inti.',
    [
        ['Lead / Inquiry Capture', 'Tidak ada', 'Sales langsung create Project status DRAFT', 'Tidak ada entity Lead/Opportunity. CRM funnel hilang.'],
        ['CRM Qualification', 'Tidak ada', '—', 'Tidak ada penilaian fit/budget sebelum project dibuat'],
        ['Services Scoping & Estimate', 'Sebagian', 'DraftCompletionCard: PM isi description/dates/revenue/mandays/cost', 'Tidak terstruktur sebagai scoping engine — tidak ada break-down per phase/role/duration'],
        ['Commercial Model & Proposal', 'Sebagian', 'Field contractValue + contractValueIncludesVat', 'Tidak ada klasifikasi T&M / Fixed Fee / Retainer / Milestone eksplisit. Tidak ada generator proposal/quote document.'],
        ['Internal Approval (high-risk)', 'Tidak ada', 'Sales/PM bisa langsung set contractValue', 'Tidak ada threshold approval untuk discount/margin rendah'],
        ['Client Acceptance / Contract Signing', 'Sebagian', 'Document upload tipe CONTRACT (base64)', 'Tidak ada formal signing workflow (e-sign, signed-at timestamp, version tracking)'],
        ['Sales-to-Delivery Handoff', 'Sebagian', 'Transisi DRAFT→OBSERVATION (Sales→PM)', 'Implicit, tidak ada meeting checklist / handoff document'],
    ],
    'Gap besar. App lompat dari "ada lead" langsung ke "project ada di sistem". Cocok kalau lead-management dipakai tool eksternal (HubSpot/Pipedrive). Kalau mau end-to-end, butuh modul CRM/Sales Pipeline.'
)

add_phase(
    'Fase II — Project Initialization & Resource Capacity',
    'Fase ini menangani inisialisasi project setelah kontrak finalized.',
    [
        ['Create Project Shell', 'Ada', 'Project entity dengan lifecycle DRAFT→OBSERVATION→ACTIVE→…', '—'],
        ['Governance Setup', 'Sebagian', 'PM assignment + Principal supervisors', 'Tidak ada penjadwalan review cadence (weekly/bi-weekly governance meetings)'],
        ['Resource Planning Engine', 'Ada', 'ProjectResource + halaman /resource-planning + Propose workflow (Principal)', '—'],
        ['Capacity Confirmed Gateway', 'Sebagian', 'Color-coding cell ≥6 destructive di Resource Planning', 'Hanya visual warning, BUKAN gateway formal. Tidak ada loop ke Re-plan/Negotiate'],
        ['Project Kickoff', 'Tidak ada', '—', 'Tidak ada formal kickoff milestone (baseline criteria, communication channels, acceptance parameters)'],
    ],
    'Cukup kuat. Resource Planning adalah fitur unggulan. Gap utama: capacity check bersifat advisory bukan blocking, dan tidak ada formal kickoff event.'
)

add_phase(
    'Fase III — Execution, Change Control & Effort Tracking',
    'Fase ini menangani delivery harian dan mekanisme feedback operasional.',
    [
        ['Project Delivery Execution', 'Ada', 'Task management + Gantt timeline + WBS + Dependencies', '—'],
        ['Scope/Commercial Change Control', 'Tidak ada', 'MGMT bisa edit contractValue/scope kapan saja', 'Tidak ada change request workflow — perubahan scope tidak punya audit, approval, atau re-estimate loop'],
        ['Time Capture (billable + non-billable)', 'Ada', 'Timesheet + Task.billable flag', '—'],
        ['Expense Capture', 'Ada', 'ProjectExpense + category', '—'],
        ['PM/Lead Approval Flow', 'Ada', 'Timesheet & Expense approval workflow', '—'],
        ['Correction Cycle', 'Sebagian', 'REJECTED status ada, tapi alur edit-resubmit tidak tegas', 'Tidak ada notifikasi otomatis ke pengaju kalau timesheet/expense ditolak'],
    ],
    'Paling matang. Gap signifikan hanya pada Change Control workflow — saat ini perubahan scope dilakukan ad-hoc tanpa formal approval atau re-estimate.'
)

add_phase(
    'Fase IV — Billing, Financial Posting & Closure',
    'Fase ini menerjemahkan delivery milestone menjadi pencatatan akuntansi.',
    [
        ['Billing Preparation (WIP, margin)', 'Sebagian', 'actualCost, actualProfit, marginPct, financials endpoint', 'WIP belum eksplisit sebagai konsep (biaya sudah keluar tapi belum invoiced)'],
        ['Draft Invoice Generation', 'Sebagian', 'BillingMilestone (manual entry % + DPP/VAT split + invoice #)', 'Tidak ada generator invoice PDF. Cuma input nomor invoice manual.'],
        ['Invoice Review Control', 'Tidak ada', 'Status INVOICED langsung valid', 'Tidak ada review approval sebelum invoice go-live'],
        ['PO Balance Tracking', 'Tidak ada', '—', 'Tidak ada entity PO (Purchase Order). Compliance issue untuk klien enterprise.'],
        ['ERP Post Sync (Xero/Accurate)', 'Tidak ada', '—', 'Tidak ada integrasi accounting. Data finance harus rekap manual.'],
        ['Project Closure Evaluation', 'Sebagian', 'Status COMPLETE/CLOSED + ADMIN_PROJECT inbox closing doc', 'Tidak ada formal post-mortem (performance log, client feedback, lesson learned)'],
    ],
    'Gap moderat-besar. Billing process terlalu manual, tidak ada PO tracking, tidak ada ERP sync.'
)

doc.add_heading('Ringkasan Strategis', level=1)

doc.add_heading('Kekuatan App Saat Ini (Fase II–III)', level=2)
for item in [
    'Resource Planning + Principal-Propose workflow',
    'WBS Gantt + Task management dengan dependencies',
    'Timesheet & Expense approval lifecycle',
    'Billing Milestone dengan VAT split + VAT Recap report',
    'Role-based access control yang granular (8+ role)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Gap Prioritas', level=2)

doc.add_heading('Priority 1 — Quick Wins (low risk, high value)', level=3)
for item in [
    'Project Kickoff Milestone — tambah status KICKOFF + checklist baseline (communication channel, acceptance criteria)',
    'Auto-notify rejection untuk timesheet/expense (email atau notifikasi in-app)',
    'Invoice Review approval — tambah status INVOICE_DRAFT → review Finance → INVOICED',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Priority 2 — Major Modules (butuh schema baru)', level=3)
for item in [
    'Change Control workflow — entity ChangeRequest dengan estimasi impact + approval → update contractValue',
    'Contract Type classification — enum: T&M / Fixed Fee / Retainer / Milestone — untuk billing rules berbeda',
    'PO Tracking — entity PurchaseOrder per project dengan balance vs invoice',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Priority 3 — Strategic (butuh integrasi/architecture)', level=3)
for item in [
    'CRM/Lead module (Fase I lengkap) — atau integrasi ke HubSpot/Pipedrive',
    'ERP integration (Xero/Accurate/Jurnal) — push invoice + pull payment status',
    'Invoice PDF generator — auto-generate dari BillingMilestone',
    'Project closure evaluation — capture lesson learned + client feedback survey',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Penutup', level=1)
doc.add_paragraph(
    "App sudah lebih dari 60% match dengan standar PSO 4-fase. Yang hilang utamanya berada di:"
)
doc.add_paragraph('Hulu (Lead → Proposal) — bisa diatasi pakai tool CRM eksternal kalau mau', style='List Bullet')
doc.add_paragraph('Hilir (Invoice review → ERP sync) — gap accounting integration', style='List Bullet')
doc.add_paragraph(
    "Rekomendasi: mulai dari Priority 1 (quick wins) untuk meningkatkan maturity tanpa risiko besar, "
    "kemudian evaluasi Priority 2 berdasarkan kebutuhan klien enterprise (terutama PO Tracking dan Change Control)."
)

import os
os.makedirs('exports', exist_ok=True)
out = 'exports/SecureProfit-Hub-Gap-Analysis.docx'
doc.save(out)
print(f'Saved: {out}')
