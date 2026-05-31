#!/usr/bin/env python3
"""Generate a detailed Xero integration plan as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x16, 0xA3, 0x4A)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = GREEN
        r.font.size = Pt(16)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = DARK
        r.font.size = Pt(13)
    return p


def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr[i], "16A34A")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


# ============ COVER ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rencana Integrasi Xero")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("SecureProfit Hub  —  Project Profitability & Billing System")
r.font.size = Pt(13)
r.font.color.rgb = GRAY

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Dokumen Perencanaan Teknis  •  Versi 1.0  •  Mei 2026")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GRAY

doc.add_paragraph()

# ============ 1. RINGKASAN ============
h1("1. Ringkasan Eksekutif")
body(
    "Dokumen ini menguraikan rencana lengkap untuk mengintegrasikan SecureProfit Hub "
    "dengan Xero (perangkat lunak akuntansi cloud). Tujuannya adalah menghilangkan input "
    "manual ganda: invoice dan data klien yang dibuat di SecureProfit Hub akan otomatis "
    "tersinkron ke Xero, dan status pembayaran dari Xero ditarik kembali ke dalam sistem."
)
body(
    "Xero tidak tersedia sebagai connector bawaan Replit, sehingga integrasi dibangun "
    "secara custom menggunakan OAuth 2.0 langsung ke Xero API. Struktur data aplikasi saat "
    "ini (BillingMilestone, InvoiceSetting, Client, perhitungan DPP/PPN 11%) sudah sangat "
    "selaras dengan model data Xero, sehingga pemetaannya relatif lurus."
)

h2("Manfaat Utama")
bullet("Eliminasi entri ganda invoice antara SecureProfit Hub dan Xero.")
bullet("Status pembayaran (PAID) otomatis tersinkron — VAT Recap & laporan keuangan selalu akurat.")
bullet("Data klien (Contact) konsisten di kedua sistem.")
bullet("Mempercepat proses tutup buku bulanan tim Finance.")

# ============ 2. RUANG LINGKUP ============
h1("2. Ruang Lingkup & Pemetaan Data")
body("Entitas yang akan disinkronkan beserta pemetaannya ke Xero:")
table(
    ["SecureProfit Hub", "Arah", "Xero", "Catatan"],
    [
        ["Client (name, contactPerson, email, phone)", "→", "Contact", "Dibuat jika belum ada"],
        ["BillingMilestone + Invoice (DPP, PPN, invoiceNumber)", "→", "Invoice (ACCREC)", "Tipe Accounts Receivable"],
        ["Project (contractValue, currency)", "→", "Invoice LineItem", "Sumber nilai & deskripsi"],
        ["PPN 11%", "→", "Tax Rate", "Dipetakan ke tax rate Xero"],
        ["Status PAID / paidAt", "←", "Invoice status & payments", "Ditarik balik ke sistem"],
    ],
    widths=[2.4, 0.5, 1.6, 1.6],
)

doc.add_paragraph()
h2("Di Luar Ruang Lingkup (Fase Awal)")
bullet("Sinkronisasi expense / bill (hutang) ke Xero.")
bullet("Rekonsiliasi bank otomatis.")
bullet("Multi-currency lanjutan (selain yang sudah didukung Project.exchangeRate).")

# ============ 3. PRASYARAT ============
h1("3. Prasyarat")
body("Hal-hal yang harus disiapkan sebelum implementasi dimulai:", bold=True)
table(
    ["Item", "Penanggung Jawab", "Keterangan"],
    [
        ["Akun Xero Developer", "Klien / Pemilik", "Daftar di developer.xero.com"],
        ["Client ID & Client Secret", "Klien / Pemilik", "Dari app yang didaftarkan di Xero"],
        ["Organisasi Xero (atau Demo Company)", "Klien / Pemilik", "Demo Company gratis untuk testing"],
        ["Domain produksi terdaftar", "Tim Dev", "Untuk redirect URI OAuth"],
        ["Penyimpanan Secret", "Tim Dev", "Replit Secrets (bukan di kode)"],
    ],
    widths=[2.3, 1.8, 2.3],
)

# ============ 4. ARSITEKTUR ============
h1("4. Arsitektur Teknis")
h2("4.1 Alur OAuth 2.0")
numbered("User (MGMT/Finance) klik \"Connect to Xero\" di halaman pengaturan.")
numbered("Backend redirect ke halaman otorisasi Xero dengan scopes yang diminta.")
numbered("User login & menyetujui akses di Xero.")
numbered("Xero redirect balik ke /api/xero/callback dengan authorization code.")
numbered("Backend menukar code menjadi access token + refresh token, lalu menyimpannya.")
numbered("Backend menyimpan tenantId (organisasi Xero yang terhubung).")

h2("4.2 Scopes yang Dibutuhkan")
bullet("offline_access — agar dapat refresh token (akses jangka panjang).")
bullet("accounting.contacts — membuat & membaca Contact.")
bullet("accounting.transactions — membuat & membaca Invoice.")
bullet("accounting.settings — membaca Tax Rate & Account.")

h2("4.3 Manajemen Token")
body(
    "Access token Xero hanya berlaku 30 menit. Sistem harus menyimpan refresh token "
    "secara aman dan otomatis memperbaruinya sebelum kedaluwarsa. Refresh token Xero "
    "berputar (rotating) — setiap kali dipakai, token baru menggantikan yang lama dan "
    "harus segera disimpan."
)

h2("4.4 Perubahan Skema Database (Prisma)")
body("Penambahan kolom/model untuk mencegah duplikasi sinkronisasi:")
table(
    ["Model", "Kolom Baru", "Tujuan"],
    [
        ["XeroConnection (baru)", "tenantId, tokens, expiresAt, connectedBy", "Menyimpan koneksi & token"],
        ["Client", "xeroContactId", "Mengikat ke Contact Xero"],
        ["BillingMilestone", "xeroInvoiceId, xeroSyncedAt", "Mengikat ke Invoice Xero"],
    ],
    widths=[2.2, 2.4, 2.2],
)

# ============ 5. RENCANA BERTAHAP ============
h1("5. Rencana Implementasi Bertahap")
body(
    "Implementasi dibagi menjadi tiga fase agar nilai bisnis tertinggi (push invoice) "
    "dapat dirilis lebih dahulu, sambil meminimalkan risiko.",
)

h2("Fase 1 — Koneksi & Push Invoice (Satu Arah)")
bullet("Daftarkan app di Xero Developer & simpan kredensial di Secrets.")
bullet("Bangun alur OAuth (connect, callback, auto-refresh token).")
bullet("Tambah skema DB (XeroConnection, xeroContactId, xeroInvoiceId).")
bullet("Sinkron Client → Xero Contact (buat jika belum ada).")
bullet("Tombol \"Push to Xero\" saat invoice di-generate → buat Invoice ACCREC.")
bullet("Halaman pengaturan koneksi (Connect/Disconnect, lihat organisasi).")
body("Hasil: invoice dari SecureProfit Hub muncul otomatis di Xero.", italic=True, color=GRAY)

h2("Fase 2 — Sinkronisasi Status Pembayaran (Dua Arah)")
bullet("Tarik status Invoice dari Xero (PAID / AUTHORISED).")
bullet("Update BillingMilestone.status & paidAt secara otomatis.")
bullet("Tombol \"Sync from Xero\" manual + opsi sinkron terjadwal.")
bullet("Refresh VAT Recap & laporan keuangan dari data terbaru.")
body("Hasil: status pembayaran selalu akurat tanpa input manual.", italic=True, color=GRAY)

h2("Fase 3 — Otomasi & Penyempurnaan (Opsional)")
bullet("Webhook Xero untuk update status pembayaran secara real-time.")
bullet("Pemetaan Tax Rate & Account Code yang dapat dikonfigurasi.")
bullet("Penanganan error & log sinkronisasi yang dapat dilihat Finance.")
bullet("Sinkron massal (bulk) untuk invoice lama.")

# ============ 6. UI ============
h1("6. Antarmuka Pengguna (UI)")
bullet("Halaman pengaturan Xero (kemungkinan di dalam /invoice-settings yang sudah ada).")
bullet("Akses dibatasi untuk peran MANAGEMENT dan FINANCE.")
bullet("Indikator status koneksi (Terhubung / Tidak terhubung + nama organisasi).")
bullet("Tombol Push to Xero pada layar invoice; tombol Sync pada layar billing.")
body(
    "Catatan: seluruh teks antarmuka tetap dalam Bahasa Inggris sesuai standar produk.",
    italic=True, color=GRAY,
)

# ============ 7. TESTING ============
h1("7. Strategi Pengujian")
bullet("Gunakan Xero Demo Company (gratis) untuk seluruh pengujian awal.")
bullet("Uji alur OAuth: connect, refresh token, disconnect, reconnect.")
bullet("Uji pembuatan Contact & Invoice — verifikasi DPP, PPN, dan total cocok.")
bullet("Uji sinkron status pembayaran dua arah.")
bullet("Uji penanganan error (token kedaluwarsa, koneksi terputus, data invalid).")
bullet("Verifikasi tidak ada duplikasi saat push berulang.")

# ============ 8. RISIKO ============
h1("8. Risiko & Mitigasi")
table(
    ["Risiko", "Dampak", "Mitigasi"],
    [
        ["Rotasi refresh token gagal disimpan", "Koneksi putus", "Transaksi atomik saat simpan token baru"],
        ["Perbedaan pembulatan DPP/PPN", "Selisih nilai invoice", "Uji ketat dengan Demo Company"],
        ["Rate limit Xero API", "Sinkron lambat", "Antrian & retry dengan backoff"],
        ["Duplikasi invoice/contact", "Data ganda", "Simpan xeroInvoiceId / xeroContactId"],
        ["Multi-organisasi (tenant)", "Salah tujuan sync", "Simpan & validasi tenantId"],
    ],
    widths=[2.2, 1.7, 2.6],
)

# ============ 9. RINGKASAN EFFORT ============
h1("9. Ringkasan Effort")
body(
    "Integrasi ini tergolong fitur skala menengah. Bagian paling menantang adalah alur "
    "OAuth (rotasi token + multi-tenant). Disarankan merilis Fase 1 (push invoice satu "
    "arah) terlebih dahulu karena memberikan nilai bisnis tertinggi, kemudian melanjutkan "
    "ke sinkronisasi status pembayaran pada Fase 2."
)
table(
    ["Fase", "Fokus", "Prioritas"],
    [
        ["Fase 1", "OAuth + Push Invoice (1 arah)", "Tinggi"],
        ["Fase 2", "Sync Status Pembayaran (2 arah)", "Menengah"],
        ["Fase 3", "Webhook & Otomasi", "Rendah / Opsional"],
    ],
    widths=[1.3, 3.6, 1.6],
)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("— Akhir Dokumen —")
r.italic = True
r.font.color.rgb = GRAY

import os
os.makedirs("exports", exist_ok=True)
out = "exports/Rencana_Integrasi_Xero_SecureProfit_Hub.docx"
doc.save(out)
print("Saved:", out)
