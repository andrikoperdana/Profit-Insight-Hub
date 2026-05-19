from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

title = doc.add_heading('Gap Analysis: SecureProfit Hub vs Standard PSO Process', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Internal Document — For Management Information')
r.italic = True
r.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    "SecureProfit Hub currently covers approximately 60% of the standard Professional Services Operations (PSO) "
    "4-phase process. Core strengths lie in Phase II (Project Initialization) and Phase III (Execution & Effort Tracking). "
    "The largest gaps are in Phase I (Lead Lifecycle / CRM) and Phase IV (Billing & ERP Integration)."
)

def add_phase(title_text, summary_text, table_data, conclusion):
    doc.add_heading(title_text, level=1)
    doc.add_paragraph(summary_text)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Status'
    hdr[2].text = 'Current Implementation'
    hdr[3].text = 'Gap'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for row_data in table_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    p = doc.add_paragraph()
    r = p.add_run('Conclusion: ')
    r.bold = True
    p.add_run(conclusion)
    doc.add_paragraph()

add_phase(
    'Phase I — Lead Lifecycle & Commercial Design',
    'This phase covers external client onboarding operations prior to core system project creation.',
    [
        ['Lead / Inquiry Capture', 'Missing', 'Sales creates Project directly in DRAFT status', 'No Lead/Opportunity entity. CRM funnel is absent.'],
        ['CRM Qualification', 'Missing', '—', 'No structural fit or budget health assessment before project creation'],
        ['Services Scoping & Estimate', 'Partial', 'DraftCompletionCard: PM fills description/dates/revenue/mandays/cost', 'Not structured as a scoping engine — no breakdown per phase/role/duration'],
        ['Commercial Model & Proposal', 'Partial', 'contractValue field + contractValueIncludesVat flag', 'No explicit classification for T&M / Fixed Fee / Retainer / Milestone. No proposal/quote document generator.'],
        ['Internal Approval (high-risk)', 'Missing', 'Sales/PM can directly set contractValue', 'No threshold-based approval for discount/margin'],
        ['Client Acceptance / Contract Signing', 'Partial', 'Document upload (CONTRACT type, base64)', 'No formal signing workflow (e-sign, signed-at timestamp, version tracking)'],
        ['Sales-to-Delivery Handoff', 'Partial', 'DRAFT→OBSERVATION transition (Sales→PM)', 'Implicit — no handoff meeting checklist or document'],
    ],
    'Significant gap. The app jumps from "lead exists" directly to "project in system". This is acceptable if lead management is handled by an external tool (HubSpot/Pipedrive). For an end-to-end solution, a CRM/Sales Pipeline module is required.'
)

add_phase(
    'Phase II — Project Initialization & Resource Capacity',
    'This phase covers project initialization after the contract is finalized.',
    [
        ['Create Project Shell', 'Present', 'Project entity with lifecycle DRAFT→OBSERVATION→ACTIVE→…', '—'],
        ['Governance Setup', 'Partial', 'PM assignment + Principal supervisors', 'No scheduled review cadence (weekly/bi-weekly governance meetings)'],
        ['Resource Planning Engine', 'Present', 'ProjectResource + /resource-planning page + Principal Propose workflow', '—'],
        ['Capacity Confirmed Gateway', 'Partial', 'Color-coded cells (≥6 destructive) on Resource Planning', 'Visual warning only — NOT a formal gateway. No loop into Re-plan/Negotiate'],
        ['Project Kickoff', 'Missing', '—', 'No formal kickoff milestone (baseline criteria, communication channels, acceptance parameters)'],
    ],
    'Reasonably strong. Resource Planning is a standout feature. Main gaps: capacity check is advisory rather than blocking, and there is no formal kickoff event.'
)

add_phase(
    'Phase III — Execution, Change Control & Effort Tracking',
    'This phase covers day-to-day delivery and operational feedback mechanisms.',
    [
        ['Project Delivery Execution', 'Present', 'Task management + Gantt timeline + WBS + Dependencies', '—'],
        ['Scope/Commercial Change Control', 'Missing', 'MGMT can edit contractValue/scope at any time', 'No change request workflow — scope changes have no audit trail, approval, or re-estimate loop'],
        ['Time Capture (billable + non-billable)', 'Present', 'Timesheet + Task.billable flag', '—'],
        ['Expense Capture', 'Present', 'ProjectExpense + category', '—'],
        ['PM/Lead Approval Flow', 'Present', 'Timesheet & Expense approval workflow', '—'],
        ['Correction Cycle', 'Partial', 'REJECTED status exists, but edit-resubmit flow is not enforced', 'No automatic notification to submitter when timesheet/expense is rejected'],
    ],
    'The most mature phase. The only significant gap is Change Control — scope changes are currently handled ad-hoc without formal approval or re-estimation.'
)

add_phase(
    'Phase IV — Billing, Financial Posting & Closure',
    'This final phase translates delivery milestones into validated accounting entries.',
    [
        ['Billing Preparation (WIP, margin)', 'Partial', 'actualCost, actualProfit, marginPct, financials endpoint', 'WIP not explicitly modeled (cost incurred but not yet invoiced)'],
        ['Draft Invoice Generation', 'Partial', 'BillingMilestone (manual entry % + DPP/VAT split + invoice #)', 'No PDF invoice generator. Invoice numbers entered manually.'],
        ['Invoice Review Control', 'Missing', 'INVOICED status is immediately valid', 'No review approval before invoice goes live'],
        ['PO Balance Tracking', 'Missing', '—', 'No Purchase Order entity. Compliance issue for enterprise clients.'],
        ['ERP Post Sync (Xero / similar)', 'Missing', '—', 'No accounting integration. Finance team must reconcile manually.'],
        ['Project Closure Evaluation', 'Partial', 'COMPLETE/CLOSED status + ADMIN_PROJECT closing-doc inbox', 'No formal post-mortem (performance log, client feedback, lessons learned)'],
    ],
    'Moderate-to-large gap. The billing process is overly manual, with no PO tracking and no ERP sync.'
)

doc.add_heading('Strategic Summary', level=1)

doc.add_heading('Current Strengths (Phases II–III)', level=2)
for item in [
    'Resource Planning + Principal-Propose workflow',
    'WBS Gantt + Task management with dependencies',
    'Timesheet & Expense approval lifecycle',
    'Billing Milestones with VAT split + VAT Recap report',
    'Granular role-based access control (8+ roles)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Priority Gaps', level=2)

doc.add_heading('Priority 1 — Quick Wins (low risk, high value)', level=3)
for item in [
    'Project Kickoff Milestone — add KICKOFF status + baseline checklist (communication channel, acceptance criteria)',
    'Auto-notify rejections for timesheet/expense (email or in-app notification)',
    'Invoice Review approval — add INVOICE_DRAFT status → Finance review → INVOICED',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Priority 2 — Major Modules (new schema required)', level=3)
for item in [
    'Change Control workflow — ChangeRequest entity with impact estimation + approval → contractValue update',
    'Contract Type classification — enum: T&M / Fixed Fee / Retainer / Milestone — to drive distinct billing rules',
    'PO Tracking — PurchaseOrder entity per project with balance vs invoice reconciliation',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Priority 3 — Strategic (requires integration/architecture)', level=3)
for item in [
    'CRM/Lead module (full Phase I) — or integration with HubSpot/Pipedrive',
    'ERP integration (Xero/Accurate/Jurnal) — push invoices + pull payment status',
    'Invoice PDF generator — auto-generate from BillingMilestone',
    'Project closure evaluation — capture lessons learned + client feedback survey',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Closing Notes', level=1)
doc.add_paragraph(
    "The application already matches more than 60% of the standard PSO 4-phase process. The principal gaps are:"
)
doc.add_paragraph('Upstream (Lead → Proposal) — can be mitigated by adopting an external CRM tool', style='List Bullet')
doc.add_paragraph('Downstream (Invoice review → ERP sync) — the accounting integration gap', style='List Bullet')
doc.add_paragraph(
    "Recommendation: start with Priority 1 (quick wins) to raise maturity at low risk, then evaluate Priority 2 "
    "based on enterprise client requirements (particularly PO Tracking and Change Control)."
)

import os
os.makedirs('exports', exist_ok=True)
out = 'exports/SecureProfit-Hub-Gap-Analysis-EN.docx'
doc.save(out)
print(f'Saved: {out}')
