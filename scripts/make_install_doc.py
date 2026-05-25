from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p

def add_bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def add_number(text):
    doc.add_paragraph(text, style="List Number")

# === Title ===
title = doc.add_heading("Panduan Instalasi SecureProfit Hub", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Full-stack Project & Profitability Management Platform")
sr.italic = True
sr.font.size = Pt(12)

doc.add_paragraph()

# === 1. Ringkasan ===
add_heading("1. Ringkasan Aplikasi", 1)
add_para(
    "SecureProfit Hub adalah aplikasi web full-stack untuk konsultan keamanan IT. "
    "Aplikasi ini mengelola siklus hidup project dari intake oleh Sales sampai delivery oleh "
    "Konsultan, sambil memantau margin profit secara real-time."
)
add_para("Stack teknologi:", bold=True)
add_bullet("Frontend: React + Vite + TypeScript + TailwindCSS + shadcn/ui")
add_bullet("Backend: Node.js + Express + Pino logger")
add_bullet("Database: PostgreSQL via Prisma ORM")
add_bullet("Auth: JWT (HS256) + bcryptjs")
add_bullet("Monorepo: pnpm workspace")

# === 2. Prasyarat ===
add_heading("2. Prasyarat Sistem", 1)
add_para("Pastikan komputer/server Anda sudah terinstal:")
add_bullet("Node.js versi 20 atau lebih baru — https://nodejs.org")
add_bullet("pnpm versi 9 atau lebih baru — install: npm install -g pnpm")
add_bullet("PostgreSQL versi 14 atau lebih baru — https://www.postgresql.org/download")
add_bullet("Git (opsional, untuk clone repository)")

# === 3. File ===
add_heading("3. File yang Dibutuhkan", 1)
add_para("Siapkan dua file berikut yang sudah Anda download:")
add_bullet("secureprofit-hub-source.tar.gz — source code (13 MB)")
add_bullet("secureprofit-hub-db.sql.gz — database dump berisi schema + data (56 KB)")

# === 4. Extract source ===
add_heading("4. Ekstrak Source Code", 1)
add_number("Buat folder kerja, contoh: C:\\Projects\\secureprofit-hub")
add_number("Pindahkan file secureprofit-hub-source.tar.gz ke folder tersebut")
add_number("Buka Command Prompt / Terminal di folder tersebut, jalankan:")
add_code("tar -xzf secureprofit-hub-source.tar.gz")
add_para(
    "Catatan: Windows 10/11 sudah punya perintah tar bawaan. "
    "Atau gunakan 7-Zip / WinRAR untuk ekstrak via klik kanan."
)

# === 5. Database ===
add_heading("5. Setup Database PostgreSQL", 1)

add_heading("5.1. Buat Database", 2)
add_para("Login ke PostgreSQL sebagai superuser (postgres), lalu jalankan:")
add_code(
    "CREATE DATABASE secureprofit;\n"
    "CREATE USER secureprofit_user WITH PASSWORD 'ganti_dengan_password_kuat';\n"
    "GRANT ALL PRIVILEGES ON DATABASE secureprofit TO secureprofit_user;"
)

add_heading("5.2. Restore Database dari Dump", 2)
add_para("Ekstrak file dump terlebih dahulu:")
add_code("gunzip secureprofit-hub-db.sql.gz")
add_para("Lalu restore ke database yang baru dibuat:")
add_code(
    'psql -U secureprofit_user -d secureprofit -h localhost -f secureprofit-hub-db.sql'
)
add_para(
    "Jika muncul prompt password, masukkan password yang Anda set di langkah 5.1. "
    "Restore berisi: schema tabel, seed users (Management, PM, Sales, Konsultan, dll), "
    "business units, skills, dan sample projects."
)

# === 6. Environment ===
add_heading("6. Konfigurasi Environment Variables", 1)
add_para("Di folder root project, buat file bernama .env dengan isi:")
add_code(
    'DATABASE_URL="postgresql://secureprofit_user:password_anda@localhost:5432/secureprofit"\n'
    'SESSION_SECRET="ganti-dengan-string-acak-panjang-minimal-32-karakter"\n'
    'NODE_ENV="production"'
)
add_para("Tips:", bold=True)
add_bullet(
    "SESSION_SECRET harus string acak yang panjang. Generate dengan: "
    "node -e \"console.log(require('crypto').randomBytes(32).toString('hex'))\""
)
add_bullet("Untuk development, set NODE_ENV=\"development\" agar tombol demo data muncul")

# === 7. Install dependencies ===
add_heading("7. Install Dependencies", 1)
add_para("Di folder root project, jalankan:")
add_code("pnpm install")
add_para(
    "Proses ini memakan waktu 2–5 menit tergantung kecepatan internet. "
    "pnpm akan mengunduh seluruh dependency untuk semua artifact di monorepo."
)

# === 8. Codegen ===
add_heading("8. Generate Code (Prisma & API Client)", 1)
add_para("Jalankan dua perintah berikut secara berurutan:")
add_code("pnpm --filter @workspace/db exec prisma generate")
add_para("Perintah di atas membuat Prisma Client untuk akses database.")
add_code("pnpm --filter @workspace/api-spec run codegen")
add_para("Perintah di atas membuat React Query hooks + Zod schemas dari OpenAPI spec.")

# === 9. Verify schema ===
add_heading("9. Verifikasi Schema Database (Opsional)", 1)
add_para(
    "Jika Anda restore dari dump terbaru, schema sudah sinkron. "
    "Untuk memastikan, jalankan:"
)
add_code("pnpm --filter @workspace/db exec prisma db push")
add_para(
    "Perintah ini akan membandingkan schema.prisma dengan database. "
    "Jika sudah sinkron, tidak ada perubahan yang diterapkan."
)

# === 10. Run ===
add_heading("10. Jalankan Aplikasi", 1)

add_heading("10.1. Mode Development", 2)
add_para("Jalankan dua workflow di dua terminal terpisah:")
add_para("Terminal 1 — API Server (port 8080):", bold=True)
add_code("pnpm --filter @workspace/api-server run dev")
add_para("Terminal 2 — Web Frontend (port 5173):", bold=True)
add_code("pnpm --filter @workspace/web run dev")
add_para("Buka browser: http://localhost:5173")

add_heading("10.2. Mode Production", 2)
add_para("Build dulu kedua artifact:")
add_code(
    "pnpm --filter @workspace/api-server run build\n"
    "pnpm --filter @workspace/web run build"
)
add_para("Lalu jalankan:")
add_code(
    "pnpm --filter @workspace/api-server run start\n"
    "pnpm --filter @workspace/web run preview"
)
add_para(
    "Untuk production sebenarnya, disarankan gunakan reverse proxy "
    "(nginx / Caddy) dengan SSL, dan process manager (PM2 / systemd) untuk API Server."
)

# === 11. Login ===
add_heading("11. Akun Login Default", 1)
add_para("Password untuk semua akun: password123", bold=True)
add_para("Daftar akun utama (domain @secureprofit.id):")

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Email"
hdr[2].text = "Nama"

accounts = [
    ("Management", "management@secureprofit.id", "Adi Wibowo"),
    ("Project Manager", "pm@secureprofit.id", "Sari Pratiwi"),
    ("Project Manager 2", "pm2@secureprofit.id", "Yusuf Maulana"),
    ("Sales", "sales@secureprofit.id", "Budi Santoso"),
    ("Konsultan", "konsultan@secureprofit.id", "Rian Hidayat"),
    ("Konsultan 2", "konsultan2@secureprofit.id", "Dewi Lestari"),
    ("Technical Writer", "writer@secureprofit.id", "Ayu Wulandari"),
    ("Admin Project", "admin@secureprofit.id", "Tono Setiawan"),
    ("Finance", "finance@secureprofit.id", "Maya Anggraini"),
    ("Site Admin", "siteadmin@secureprofit.id", "Rina Kartika"),
    ("HR", "hr@itsecasia.com", "Sinta Permata"),
    ("Principal Konsultan", "principal.kon.h7q4@itsecasia.com", "Bayu Prasetyo"),
    ("Principal Tech Writer", "principal.tw.m9k2@itsecasia.com", "Indah Kusumawardani"),
    ("Principal Admin Project", "principal.ap.r3n8@itsecasia.com", "Fajar Nugroho"),
]
for role, email, name in accounts:
    row = table.add_row().cells
    row[0].text = role
    row[1].text = email
    row[2].text = name

doc.add_paragraph()
add_para(
    "PENTING: Setelah aplikasi live di production, segera ganti semua password default "
    "lewat menu Settings.",
    bold=True,
)

# === 12. Troubleshooting ===
add_heading("12. Troubleshooting", 1)

add_para("Masalah: 'Cannot connect to database'", bold=True)
add_bullet("Pastikan PostgreSQL service berjalan")
add_bullet("Cek DATABASE_URL di file .env — username, password, host, port, nama database")
add_bullet("Test koneksi manual: psql -U secureprofit_user -d secureprofit -h localhost")

add_para("Masalah: 'pnpm: command not found'", bold=True)
add_bullet("Install pnpm: npm install -g pnpm")
add_bullet("Restart terminal setelah install")

add_para("Masalah: Port sudah digunakan", bold=True)
add_bullet("Ubah variable PORT di .env, contoh: PORT=8090")
add_bullet("Atau kill process yang pakai port tersebut")

add_para("Masalah: Tipe data error setelah pull update", bold=True)
add_bullet("Ulangi langkah 8 — generate Prisma client dan API codegen")
add_bullet("Jalankan: pnpm install untuk sync dependency")

add_para("Masalah: Halaman blank / error di browser", bold=True)
add_bullet("Buka DevTools (F12) → tab Console, lihat error message")
add_bullet("Pastikan API Server berjalan di port yang benar")
add_bullet("Cek file .env tidak ada typo")

# === 13. Backup ===
add_heading("13. Backup Database Berkala", 1)
add_para("Untuk backup database production, jalankan secara rutin:")
add_code(
    'pg_dump -U secureprofit_user -d secureprofit -h localhost --no-owner --no-acl '
    '-f backup_$(date +%Y%m%d).sql'
)
add_para("Disarankan setup cron job harian dan simpan minimal 7 hari ke belakang.")

# === 14. Support ===
add_heading("14. Dukungan & Dokumentasi", 1)
add_bullet("File replit.md di root project berisi overview teknis lengkap")
add_bullet("Folder lib/db/prisma/schema.prisma berisi struktur database")
add_bullet("Folder lib/api-spec/openapi.yaml berisi spesifikasi API")
add_bullet("Reseed database: pnpm --filter @workspace/db run seed")

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("— Akhir Dokumen —")
fr.italic = True
fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save("dist/Panduan-Instalasi-SecureProfit-Hub.docx")
print("OK")
